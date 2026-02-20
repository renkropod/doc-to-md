"""DOCX 파서 - python-docx + mammoth"""

import logging
from pathlib import Path

from doc_to_md.exceptions import ParserError
from doc_to_md.parsers.base import BaseParser, ParseResult
from doc_to_md.utils.image_handler import generate_image_filename, save_image

logger = logging.getLogger(__name__)

# python-docx 스타일 → 마크다운 헤딩 매핑
HEADING_MAP = {
    "Heading 1": "# ",
    "Heading 2": "## ",
    "Heading 3": "### ",
    "Heading 4": "#### ",
    "Heading 5": "##### ",
    "Heading 6": "###### ",
    "Title": "# ",
    "Subtitle": "## ",
}


class DocxParser(BaseParser):
    """DOCX 문서를 마크다운으로 변환하는 파서"""

    def supported_extensions(self) -> list[str]:
        return [".docx", ".doc"]

    def parse(self, file_path: Path, **kwargs) -> ParseResult:
        """DOCX 파일을 파싱

        mammoth을 1순위로 시도하고, 실패 시 python-docx로 fallback합니다.
        """
        extract_images = kwargs.get("extract_images", False)
        image_dir = kwargs.get("image_dir")

        try:
            return self._parse_with_mammoth(file_path, extract_images, image_dir)
        except Exception as e:
            logger.warning("mammoth 파싱 실패, python-docx로 fallback: %s", e)

        try:
            return self._parse_with_python_docx(file_path, extract_images, image_dir)
        except Exception as e:
            raise ParserError(f"DOCX 파싱 실패: {file_path} - {e}") from e

    def _parse_with_mammoth(
        self,
        file_path: Path,
        extract_images: bool,
        image_dir: Path | None,
    ) -> ParseResult:
        """mammoth을 사용한 DOCX→HTML→Markdown 변환"""
        import mammoth
        import markdownify

        images: list[dict] = []
        img_index = 0

        def convert_image(image):
            nonlocal img_index
            with image.open() as img_stream:
                img_data = img_stream.read()
                content_type = image.content_type or "image/png"
                ext = content_type.split("/")[-1]
                if ext == "jpeg":
                    ext = "jpg"

                if extract_images and image_dir:
                    filename = generate_image_filename(
                        file_path.stem, img_index, ext
                    )
                    save_image(img_data, image_dir, filename)
                    rel_path = f"./{image_dir.name}/{filename}"
                    images.append({
                        "path": str(image_dir / filename),
                        "alt": f"Image {img_index}",
                        "page": 0,
                    })
                    img_index += 1
                    return {"src": rel_path}

                img_index += 1
                return {}

        with open(file_path, "rb") as f:
            result = mammoth.convert_to_html(
                f,
                convert_image=mammoth.images.img_element(convert_image),
            )

        for msg in result.messages:
            logger.debug("mammoth 메시지: %s", msg)

        md_content = markdownify.markdownify(
            result.value,
            heading_style="ATX",
            strip=["script", "style"],
        )

        return ParseResult(
            content=md_content.strip(),
            images=images,
        )

    def _parse_with_python_docx(
        self,
        file_path: Path,
        extract_images: bool,
        image_dir: Path | None,
    ) -> ParseResult:
        """python-docx를 사용한 DOCX 파싱"""
        from docx import Document
        from docx.table import Table as DocxTable

        doc = Document(str(file_path))
        parts: list[str] = []
        tables: list[str] = []
        images: list[dict] = []
        img_index = 0

        metadata = self._extract_metadata(doc)

        if extract_images and image_dir:
            images, img_index = self._extract_images(doc, image_dir, file_path.stem)

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                para = self._find_paragraph(doc, element)
                if para:
                    md_line = self._paragraph_to_markdown(para)
                    if md_line:
                        parts.append(md_line)

            elif tag == "tbl":
                table = self._find_table(doc, element)
                if table:
                    md_table = self._table_to_markdown(table)
                    tables.append(md_table)
                    parts.append(md_table)

        return ParseResult(
            content="\n\n".join(parts),
            metadata=metadata,
            images=images,
            tables=tables,
        )

    def _find_paragraph(self, doc, element):
        """element에 해당하는 Paragraph 객체 찾기"""
        for para in doc.paragraphs:
            if para._element is element:
                return para
        return None

    def _find_table(self, doc, element):
        """element에 해당하는 Table 객체 찾기"""
        for table in doc.tables:
            if table._element is element:
                return table
        return None

    def _paragraph_to_markdown(self, paragraph) -> str:
        """Paragraph를 마크다운 문자열로 변환"""
        text = paragraph.text.strip()
        if not text:
            return ""

        style_name = paragraph.style.name if paragraph.style else ""

        prefix = HEADING_MAP.get(style_name, "")
        if prefix:
            return f"{prefix}{text}"

        if style_name.startswith("List"):
            return f"- {text}"

        runs = paragraph.runs
        if runs:
            parts = []
            for run in runs:
                t = run.text
                if not t:
                    continue
                if run.bold and run.italic:
                    t = f"***{t}***"
                elif run.bold:
                    t = f"**{t}**"
                elif run.italic:
                    t = f"*{t}*"
                parts.append(t)
            return "".join(parts)

        return text

    def _table_to_markdown(self, table) -> str:
        """docx Table을 마크다운 테이블로 변환"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            rows.append(cells)

        if not rows:
            return ""

        col_count = max(len(r) for r in rows)
        lines = [
            "| " + " | ".join(rows[0] + [""] * (col_count - len(rows[0]))) + " |",
            "| " + " | ".join("---" for _ in range(col_count)) + " |",
        ]

        for row in rows[1:]:
            padded = row + [""] * (col_count - len(row))
            lines.append("| " + " | ".join(padded[:col_count]) + " |")

        return "\n".join(lines)

    def _extract_metadata(self, doc) -> dict:
        """DOCX 메타데이터 추출"""
        props = doc.core_properties
        metadata = {}
        if props.title:
            metadata["title"] = props.title
        if props.author:
            metadata["author"] = props.author
        if props.created:
            metadata["created"] = str(props.created)
        if props.modified:
            metadata["modified"] = str(props.modified)
        if props.subject:
            metadata["subject"] = props.subject
        return metadata

    def _extract_images(
        self,
        doc,
        image_dir: Path,
        base_name: str,
    ) -> tuple[list[dict], int]:
        """DOCX에서 이미지 추출"""
        images: list[dict] = []
        img_index = 0

        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    content_type = rel.target_part.content_type
                    ext = content_type.split("/")[-1] if content_type else "png"
                    if ext == "jpeg":
                        ext = "jpg"

                    filename = generate_image_filename(base_name, img_index, ext)
                    save_image(image_data, image_dir, filename)
                    images.append({
                        "path": str(image_dir / filename),
                        "alt": f"Image {img_index}",
                        "page": 0,
                    })
                    img_index += 1
                except Exception as e:
                    logger.debug("이미지 추출 실패: %s", e)

        return images, img_index
